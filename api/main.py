from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
import shutil
import os
import uuid
import json

from agents.xml_parser import parse_jira_xml
from agents.github_scanner import scan_github_repo
from agents.module_detector import detect_module_and_channel, get_module_description
from agents.impact_engine import analyze_impact
from agents.generator_engine import generate_test_cases
from agents.reviewer_engine import review_coverage
from agents.config_loader import load_config
from ml.risk_predictor import calculate_risk_score
from agents.governance_engine import apply_governance_rules, get_risk_color

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from xml.sax.saxutils import escape as xml_escape
from docx import Document
from openpyxl import Workbook


app = FastAPI()
templates = Jinja2Templates(directory="templates")

# Store latest result (in-memory for dashboard)
latest_result = {"test_cases": [], "risk_score": None}


# ============================
# HEALTH CHECK
# ============================
@app.get("/health")
def health_check():
    return {"status": "ok", "version": "2.0.0"}


# ============================
# GET DASHBOARD
# ============================
@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard_page(request: Request):
    return templates.TemplateResponse("dashboard.html", {
        "request": request,
        **latest_result
    })


# ============================
# POST → RUN ENGINE
# ============================
@app.post("/dashboard")
async def analyze_dashboard(
    request: Request,
    file: UploadFile = File(...),
    repo_url: str = Form(""),
):
    global latest_result

    os.makedirs("reports", exist_ok=True)
    os.makedirs("temp_uploads", exist_ok=True)

    temp_path = None

    try:
        # Save uploaded file with a unique name to avoid collisions
        unique_name = f"{uuid.uuid4()}.xml"
        temp_path = os.path.join("temp_uploads", unique_name)

        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # ── Parse story ──────────────────────────────────────────
        story = parse_jira_xml(temp_path)
        config = load_config("config.yaml")
        impact = analyze_impact(story)

        # ── Module / channel detection ───────────────────────────
        module_name, channel_name = detect_module_and_channel(story)

        # ── GitHub code scan (optional) ──────────────────────────
        repo_info = None
        if repo_url and repo_url.strip():
            github_result = scan_github_repo(repo_url.strip())
            repo_info = {
                "repo_name": github_result.get("repo_name"),
                "total_files_scanned": github_result.get("total_files_scanned", 0),
                "scanned_files": github_result.get("scanned_files", []),
                "error": github_result.get("error"),
            }
            raw_signals = github_result.get("risk_signals", {})
            # Map generic signals to the format expected by the ML risk predictor
            code_risk = {
                "validations":    raw_signals.get("validations", 0),
                "fee_logic":      raw_signals.get("payment_logic", 0),
                "limit_checks":   raw_signals.get("limit_checks", 0),
                "error_handling": raw_signals.get("error_handling", 0),
            }
        else:
            code_risk = {"validations": 0, "fee_logic": 0, "limit_checks": 0, "error_handling": 0}

        # ── Risk scoring ─────────────────────────────────────────
        risk_score, risk_level, confidence = calculate_risk_score(impact, story, code_risk)
        risk_color = get_risk_color(risk_score)

        # ── Test case generation ─────────────────────────────────
        test_cases = generate_test_cases(
            story,
            impact,
            config,
            code_risk,
            risk_score,
            risk_level,
        )

        # ── Coverage review ──────────────────────────────────────
        review = review_coverage(impact, test_cases, risk_level)

        # ── Governance ───────────────────────────────────────────
        governance = apply_governance_rules(
            risk_score=risk_score,
            coverage_score=review["coverage_score"],
            environment=config.get("environment"),
        )

        # ── Build result ─────────────────────────────────────────
        latest_result = {
            "engine_version": "2.0.0",
            "channel": channel_name,
            "environment": config.get("environment"),
            "story_summary": story["summary"],
            "priority": story.get("priority", "Medium"),
            "word_count": story["word_count"],
            "impact_areas": impact,
            "risk_score": risk_score,
            "risk_level": risk_level,
            "risk_color": risk_color,
            "confidence": confidence,
            "coverage_score": review["coverage_score"],
            "missing_areas": review["missing_areas"],
            "explanation": review["explanation"],
            "recommendations": review.get("recommendations", []),
            "review_status": review.get("review_status", "FAIL"),
            "governance_decision": governance.get("decision", "REVIEW"),
            "test_cases": test_cases,
            "repo_info": repo_info,
            "repo_url": repo_url.strip() if repo_url else None,
        }

        with open("reports/execution_report.json", "w") as f:
            json.dump(latest_result, f, indent=4)

        return templates.TemplateResponse("dashboard.html", {
            "request": request,
            **latest_result
        })

    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)


# ============================
# EXPORT PDF
# ============================
@app.get("/export/pdf")
def export_pdf():
    if not latest_result.get("test_cases"):
        return {"error": "No test cases available"}

    os.makedirs("reports", exist_ok=True)
    file_path = "reports/test_cases.pdf"

    doc = SimpleDocTemplate(file_path)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("<b>Generated Test Cases</b>", styles["Heading1"]))
    elements.append(Spacer(1, 0.3 * inch))

    for tc in latest_result["test_cases"]:
        steps_html = xml_escape(tc['steps']).replace('\n', '<br/>')
        content = (
            f"<b>{xml_escape(tc['test_case_id'])} - {xml_escape(tc['summary'])}</b><br/>"
            f"Module: {xml_escape(tc['module'])} | Channel: {xml_escape(tc['channel'])} | "
            f"Priority: {xml_escape(tc['priority'])} | Type: {xml_escape(tc['test_type'])}<br/>"
            f"<b>Precondition:</b> {xml_escape(tc['precondition'])}<br/>"
            f"<b>Steps:</b><br/>{steps_html}<br/>"
            f"<b>Expected:</b> {xml_escape(tc['expected_result'])}<br/><br/>"
        )
        elements.append(Paragraph(content, styles["Normal"]))
        elements.append(Spacer(1, 0.2 * inch))

    doc.build(elements)

    return FileResponse(file_path, media_type="application/pdf", filename="test_cases.pdf")


# ============================
# EXPORT DOCX
# ============================
@app.get("/export/docx")
def export_docx():
    if not latest_result.get("test_cases"):
        return {"error": "No test cases available"}

    os.makedirs("reports", exist_ok=True)
    file_path = "reports/test_cases.docx"

    document = Document()
    document.add_heading("Generated Test Cases", level=1)

    for tc in latest_result["test_cases"]:
        document.add_heading(f"{tc['test_case_id']} - {tc['summary']}", level=2)
        document.add_paragraph(f"Module: {tc['module']}")
        document.add_paragraph(f"Channel: {tc['channel']}")
        document.add_paragraph(f"Priority: {tc['priority']}")
        document.add_paragraph(f"Type: {tc['test_type']}")
        document.add_paragraph("Precondition:")
        document.add_paragraph(tc["precondition"])
        document.add_paragraph("Steps:")
        document.add_paragraph(tc["steps"])
        document.add_paragraph("Expected Result:")
        document.add_paragraph(tc["expected_result"])
        document.add_page_break()

    document.save(file_path)

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="test_cases.docx",
    )


# ============================
# EXPORT EXCEL
# ============================
@app.get("/export/excel")
def export_excel():
    if not latest_result.get("test_cases"):
        return {"error": "No test cases available"}

    os.makedirs("reports", exist_ok=True)
    file_path = "reports/test_cases.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Test Cases"

    headers = [
        "Test Case ID", "Module", "Channel", "Summary",
        "Priority", "Test Type", "Precondition", "Steps", "Expected Result",
    ]
    ws.append(headers)

    for tc in latest_result["test_cases"]:
        ws.append([
            tc["test_case_id"],
            tc["module"],
            tc["channel"],
            tc["summary"],
            tc["priority"],
            tc["test_type"],
            tc["precondition"],
            tc["steps"],
            tc["expected_result"],
        ])

    wb.save(file_path)

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="test_cases.xlsx",
    )


# ============================
# AI CHAT ENDPOINT
# ============================
def _fallback_chat(message: str, context: dict) -> str:
    """Rule-based fallback when Claude API key is not set."""
    msg = message.lower()

    if not context.get("risk_score"):
        return ("No analysis has been run yet. Upload a user story XML file first, "
                "then I can help you understand the results.")

    # GitHub / API errors
    if any(k in msg for k in ["github", "repo", "403", "api error", "api not working",
                               "rate limit", "scan", "why", "error"]):
        repo = context.get("repo_info")
        if repo and repo.get("error"):
            err = repo["error"]
            if "403" in err or "rate limit" in err.lower():
                return ("**GitHub API Rate Limit Hit (403)**\n\n"
                        "GitHub allows only **60 unauthenticated API requests/hour** per IP.\n\n"
                        "**How to fix:**\n"
                        "- Wait ~60 minutes, then re-analyze\n"
                        "- Create a GitHub Personal Access Token (github.com → Settings → "
                        "Developer Settings → Personal Access Tokens) and add it to the config")
            return f"**GitHub scan error:** {err}"
        if repo and repo.get("total_files_scanned", 0) > 0:
            return (f"GitHub scan succeeded: **{repo['repo_name']}** — "
                    f"{repo['total_files_scanned']} files scanned for risk signals.")
        return "No GitHub repo was scanned in this run, or the scan returned no files."

    # Risk
    if any(k in msg for k in ["risk", "score", "level", "danger", "critical", "high"]):
        areas = context.get("impact_areas", [])
        return (f"**Risk Score: {context.get('risk_score')}/100 — {context.get('risk_level')}**\n\n"
                f"Driven by **{len(areas)} impact areas:**\n"
                + "\n".join(f"- {a}" for a in areas) +
                f"\n\nML Confidence: {context.get('confidence')}%")

    # Coverage / missing
    if any(k in msg for k in ["coverage", "missing", "gap", "uncovered"]):
        missing = context.get("missing_areas", [])
        score = context.get("coverage_score")
        if missing:
            return (f"Coverage: **{score}%** — gaps found:\n"
                    + "\n".join(f"- {m}" for m in missing)
                    + "\n\nAdd test cases targeting these areas to reach 100%.")
        return f"Coverage: **{score}%** — all impact areas are covered. ✓"

    # Test cases
    if any(k in msg for k in ["test case", "testcase", "how many", "generated", "tc"]):
        tcs = context.get("test_cases", [])
        priorities: dict = {}
        for tc in tcs:
            p = tc.get("priority", "Unknown")
            priorities[p] = priorities.get(p, 0) + 1
        breakdown = " | ".join(f"{k}: {v}" for k, v in priorities.items())
        return (f"**{len(tcs)} test cases generated**\n\n"
                f"Priority breakdown: {breakdown}\n"
                f"Module: {tcs[0].get('module', 'N/A') if tcs else 'N/A'} | "
                f"Channel: {tcs[0].get('channel', 'N/A') if tcs else 'N/A'}")

    # Impact areas
    if any(k in msg for k in ["impact", "area", "affect"]):
        areas = context.get("impact_areas", [])
        return ("**Detected Impact Areas:**\n"
                + "\n".join(f"- {a}" for a in areas))

    # Governance
    if any(k in msg for k in ["governance", "decision", "approve", "block", "deploy"]):
        decision = context.get("governance_decision", "N/A")
        missing = context.get("missing_areas", [])
        out = f"**Governance Decision: {decision}**\n\n"
        if decision == "APPROVE":
            out += "Risk and coverage thresholds met. Safe to proceed to deployment."
        elif decision == "REVIEW":
            out += "Manual QA review recommended before deployment."
            if missing:
                out += f"\nFix missing coverage: {', '.join(missing)}"
        elif decision == "BLOCK":
            out += "Blocked — risk score too high or coverage too low for safe deployment."
        return out

    # Story
    if any(k in msg for k in ["story", "summary", "feature", "about"]):
        return (f"**Story:** {context.get('story_summary', 'N/A')}\n"
                f"**Priority:** {context.get('priority', 'N/A')} | "
                f"**Word Count:** {context.get('word_count', 'N/A')} | "
                f"**Environment:** {context.get('environment', 'N/A')}")

    # Suggestions
    if any(k in msg for k in ["suggest", "more test", "improve", "recommend"]):
        recs = context.get("recommendations", [])
        missing = context.get("missing_areas", [])
        if recs:
            return "**Recommendations:**\n" + "\n".join(f"- {r}" for r in recs[:5])
        if missing:
            return ("**Suggested improvements:**\n"
                    + "\n".join(f"- Add test cases for: {m}" for m in missing))
        return "Coverage is complete. Consider adding edge cases, load tests, and security tests."

    # General QA knowledge (works even without an active analysis)
    if any(k in msg for k in ["boundary", "bva", "boundary value"]):
        return ("**Boundary Value Analysis (BVA)**\n\n"
                "Test at the edges of valid input ranges:\n"
                "- **Below minimum** — expect rejection\n"
                "- **At minimum** — expect acceptance\n"
                "- **Just above minimum** — expect acceptance\n"
                "- **Just below maximum** — expect acceptance\n"
                "- **At maximum** — expect acceptance\n"
                "- **Above maximum** — expect rejection\n\n"
                "BVA catches the most defects with the fewest test cases.")

    if any(k in msg for k in ["equivalence", "partition", "epa", "ep "]):
        return ("**Equivalence Partitioning**\n\n"
                "Divide inputs into groups that should behave identically:\n"
                "- **Valid partition** — one representative test per group\n"
                "- **Invalid partition** — one representative per invalid group\n\n"
                "Reduces the total number of test cases while maintaining coverage. "
                "Combine with BVA for maximum effectiveness.")

    if any(k in msg for k in ["smoke", "sanity", "regression"]):
        return ("**Key Test Types:**\n\n"
                "- **Smoke Test** — Quick build verification (10–15 min), core flows only\n"
                "- **Sanity Test** — Focused check after a bug fix or small change\n"
                "- **Regression Test** — Full suite to ensure nothing broke after changes\n"
                "- **Exploratory Test** — Unscripted, experience-based investigation\n\n"
                "Run smoke first → if pass, run full regression in CI/CD pipeline.")

    if any(k in msg for k in ["agile", "sprint", "scrum", "shift left", "shift-left"]):
        return ("**Agile QA Best Practices:**\n\n"
                "- **Shift-Left** — Test requirements and designs, not just code\n"
                "- **Definition of Done** — Include passing tests as a DoD criterion\n"
                "- **3 Amigos** — BA + Dev + QA review stories before development\n"
                "- **Continuous Testing** — Tests run on every commit via CI/CD\n"
                "- **Test-First** — Write acceptance criteria before coding begins\n"
                "- **Sprint Reviews** — Demo with actual test evidence, not just features")

    if any(k in msg for k in ["risk based", "risk-based", "prioritise", "prioritize"]):
        return ("**Risk-Based Testing:**\n\n"
                "Focus effort where failure impact is highest:\n"
                "1. **Identify risk** — What breaks if this fails? Financial loss? Data loss?\n"
                "2. **Assess likelihood** — How often does this area change? Complex code?\n"
                "3. **Prioritise** — Risk = Impact × Likelihood\n"
                "4. **Allocate tests** — Most critical paths get Critical priority TCs\n\n"
                "This engine uses ML + code signals to do this automatically for your stories.")

    if any(k in msg for k in ["best practice", "strategy", "how to test", "good test"]):
        return ("**QA Best Practices:**\n\n"
                "- **Test early** — Catch defects in requirements, not production\n"
                "- **Independent testing** — QE should not test their own code\n"
                "- **Traceability** — Each TC maps to a requirement or user story\n"
                "- **Test data management** — Use realistic, anonymised test data\n"
                "- **Defect clustering** — 80% of bugs live in 20% of modules (Pareto)\n"
                "- **Exit criteria** — Define PASS/FAIL thresholds before testing starts\n"
                "- **Automation first for** — Regression, smoke, data-driven tests\n"
                "- **Manual first for** — Exploratory, usability, new feature validation")

    # Default — full summary
    tcs = context.get("test_cases", [])
    repo = context.get("repo_info")
    repo_line = ""
    if repo:
        if repo.get("error"):
            repo_line = f"\n- GitHub Scan: ⚠ {repo['error'][:60]}"
        else:
            repo_line = f"\n- GitHub Scan: ✓ {repo.get('repo_name')} ({repo.get('total_files_scanned', 0)} files)"
    return (f"**Current Analysis Summary:**\n"
            f"- Story: {str(context.get('story_summary', 'N/A'))[:70]}\n"
            f"- Risk: **{context.get('risk_score')}/100** ({context.get('risk_level')})\n"
            f"- Test Cases: **{len(tcs)}** generated\n"
            f"- Coverage: **{context.get('coverage_score')}%**\n"
            f"- Governance: **{context.get('governance_decision', 'N/A')}**"
            f"{repo_line}\n\n"
            "Ask me about risk, coverage gaps, test cases, governance, or the GitHub scan!")


@app.post("/chat")
async def chat_endpoint(request: Request):
    data = await request.json()
    user_message = data.get("message", "").strip()
    if not user_message:
        return {"response": "Please ask a question."}

    ctx = latest_result
    ctx_parts = []
    if ctx.get("story_summary"):
        ctx_parts.append(f"Story: {ctx['story_summary']} (Priority: {ctx.get('priority', 'N/A')})")
    if ctx.get("risk_score") is not None:
        ctx_parts.append(f"Risk Score: {ctx['risk_score']}/100 — {ctx.get('risk_level', 'N/A')}")
    if ctx.get("coverage_score") is not None:
        ctx_parts.append(f"Test Coverage: {ctx['coverage_score']}% ({ctx.get('review_status', 'N/A')})")
    if ctx.get("test_cases"):
        ctx_parts.append(f"Test Cases Generated: {len(ctx['test_cases'])}")
    if ctx.get("impact_areas"):
        ctx_parts.append(f"Impact Areas: {', '.join(ctx['impact_areas'])}")
    if ctx.get("missing_areas"):
        ctx_parts.append(f"Missing Coverage: {', '.join(ctx['missing_areas'])}")
    if ctx.get("governance_decision"):
        ctx_parts.append(f"Governance: {ctx['governance_decision']}")
    if ctx.get("confidence"):
        ctx_parts.append(f"ML Confidence: {ctx['confidence']}%")
    if ctx.get("recommendations"):
        ctx_parts.append(f"Recommendations: {'; '.join(ctx['recommendations'][:3])}")
    repo = ctx.get("repo_info")
    if repo:
        if repo.get("error"):
            ctx_parts.append(f"GitHub Scan Error: {repo['error']}")
        elif repo.get("repo_name"):
            ctx_parts.append(f"GitHub Scan: {repo['repo_name']} — {repo.get('total_files_scanned', 0)} files scanned")
    context_str = "\n".join(ctx_parts) if ctx_parts else "No analysis run yet."

    try:
        import anthropic as _anthropic
        client = _anthropic.Anthropic()
        system_prompt = (
            "You are an expert QA Engineer and AI-powered test intelligence assistant. "
            "You have deep knowledge of software testing, risk assessment, QA best practices, "
            "test design techniques (BVA, equivalence partitioning, decision tables, state transition), "
            "security testing, performance testing, compliance (PCI, AML, GDPR, OWASP), "
            "CI/CD integration, agile QA, shift-left testing, and test management.\n\n"
            f"Current analysis context:\n{context_str}\n\n"
            "Answer questions about BOTH:\n"
            "1. The current analysis — risk score, coverage gaps, test cases, governance, impact areas\n"
            "2. General QA and software testing topics — techniques, best practices, strategies, tools\n"
            "Be concise and specific. Use bullet points for lists. "
            "Format key values in **bold**. For general questions without analysis context, give expert guidance."
        )
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=512,
            system=system_prompt,
            messages=[{"role": "user", "content": user_message}],
        )
        return {"response": resp.content[0].text}
    except Exception:
        return {"response": _fallback_chat(user_message, ctx)}
